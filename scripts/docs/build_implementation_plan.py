from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[2]
OUT = ROOT / "docs" / "implementation" / "VIDSLOOM_Implementation_Plan.docx"
LOGO = ROOT / "assets" / "brand" / "VIDSLOOM_Logo.png"

BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 77, 120)
INK = RGBColor(18, 32, 52)
MUTED = RGBColor(90, 100, 115)
LIGHT_GRAY = "F2F4F7"
MID_GRAY = "DADCE0"
PURPLE = RGBColor(124, 58, 237)


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in [("top", top), ("start", start), ("bottom", bottom), ("end", end)]:
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_cell_border(cell, color=MID_GRAY, size="4"):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right"):
        tag = f"w:{edge}"
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_run_font(run, name="Calibri", size=None, color=None, bold=None, italic=None):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = color
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def paragraph_text(paragraph, text, bold=False, italic=False, color=None):
    run = paragraph.add_run(text)
    set_run_font(run, color=color, bold=bold, italic=italic)
    return run


def configure_styles(doc):
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor(0, 0, 0)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    for style_name, size, color, before, after in [
        ("Heading 1", 16, BLUE, 16, 8),
        ("Heading 2", 13, BLUE, 12, 6),
        ("Heading 3", 12, DARK_BLUE, 8, 4),
    ]:
        style = styles[style_name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.color.rgb = color
        style.font.bold = True
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    for list_style_name in ["List Bullet", "List Number"]:
        style = styles[list_style_name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(11)
        style.paragraph_format.space_after = Pt(8)
        style.paragraph_format.line_spacing = 1.167


def configure_section(section):
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)
    header = section.header.paragraphs[0]
    header.text = "VIDSLOOM Implementation Plan"
    header.alignment = WD_ALIGN_PARAGRAPH.LEFT
    header.runs[0].font.size = Pt(9)
    header.runs[0].font.color.rgb = MUTED
    footer = section.footer.paragraphs[0]
    footer.text = "Build with Gemini XPRIZE | Small Business Services | Draft for execution"
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer.runs[0].font.size = Pt(9)
    footer.runs[0].font.color.rgb = MUTED


def add_title_page(doc):
    if LOGO.exists():
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(str(LOGO), width=Inches(1.15))
        p.paragraph_format.space_after = Pt(10)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run("VIDSLOOM")
    set_run_font(run, size=26, color=PURPLE, bold=True)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run("Start-to-Finish Implementation Plan")
    set_run_font(run, size=18, color=INK, bold=True)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(16)
    run = p.add_run("Build with Gemini XPRIZE business launch plan through August 17, 2026")
    set_run_font(run, size=11, color=MUTED)

    rows = [
        ("Competition category", "Small Business Services"),
        ("Submission deadline", "August 17, 2026 at 1:00 PM PDT / August 18, 2026 at 4:00 AM SGT"),
        ("Plan date", "June 13, 2026"),
        ("Staging app", "https://vidsloom-staging-feiefrrvkq-as.a.run.app"),
        ("Production app", "https://vidsloom-production-feiefrrvkq-as.a.run.app"),
        ("Primary objective", "Launch VIDSLOOM as a real paid business with production AI operations, customer evidence, and revenue proof before deadline."),
        ("Operating thesis", "Sell a service-assisted AI content operations MVP now, with automation settings, trend intelligence, approval workflows, and OAuth-gated publishing that can deepen each week."),
    ]
    add_kv_table(doc, rows, [1.75, 4.65])

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(8)
    paragraph_text(
        p,
        "Executive instruction: ",
        bold=True,
        color=INK,
    )
    paragraph_text(
        p,
        "Do not wait for full autonomous multi-platform publishing before selling. The deadline rewards proof of a viable business, AI-native operations, and category impact. Build the workflow needed to collect payment, onboard customers, generate strong campaign packs, queue/schedule posts with approval controls, log AI decisions, and produce credible evidence.",
    )
    doc.add_page_break()


def add_heading(doc, text, level=1):
    return doc.add_heading(text, level=level)


def add_para(doc, text="", bold_prefix=None):
    p = doc.add_paragraph()
    if bold_prefix and text.startswith(bold_prefix):
        paragraph_text(p, bold_prefix, bold=True, color=INK)
        paragraph_text(p, text[len(bold_prefix):])
    else:
        paragraph_text(p, text)
    return p


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        paragraph_text(p, item)


def add_numbered(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Number")
        paragraph_text(p, item)


def set_table_geometry(table, widths):
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    total_twips = sum(int(w * 1440) for w in widths)
    tbl_w.set(qn("w:w"), str(total_twips))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_layout = tbl_pr.find(qn("w:tblLayout"))
    if tbl_layout is None:
        tbl_layout = OxmlElement("w:tblLayout")
        tbl_pr.append(tbl_layout)
    tbl_layout.set(qn("w:type"), "fixed")

    existing_grid = tbl.tblGrid
    if existing_grid is not None:
        tbl.remove(existing_grid)
    grid = OxmlElement("w:tblGrid")
    for w in widths:
        grid_col = OxmlElement("w:gridCol")
        grid_col.set(qn("w:w"), str(int(w * 1440)))
        grid.append(grid_col)
    tbl.insert(1, grid)

    for row in table.rows:
        for i, cell in enumerate(row.cells):
            width = Inches(widths[i])
            cell.width = width
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(int(widths[i] * 1440)))
            tc_w.set(qn("w:type"), "dxa")
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            set_cell_margins(cell)
            set_cell_border(cell)


def style_table(table, widths, header=True):
    set_table_geometry(table, widths)
    for row_index, row in enumerate(table.rows):
        if header and row_index == 0:
            tr_pr = row._tr.get_or_add_trPr()
            tbl_header = tr_pr.find(qn("w:tblHeader"))
            if tbl_header is None:
                tbl_header = OxmlElement("w:tblHeader")
                tr_pr.append(tbl_header)
            tbl_header.set(qn("w:val"), "true")
        for cell in row.cells:
            for p in cell.paragraphs:
                p.paragraph_format.space_before = Pt(0)
                p.paragraph_format.space_after = Pt(2)
                for run in p.runs:
                    set_run_font(run, size=9.3)
            if header and row_index == 0:
                set_cell_shading(cell, LIGHT_GRAY)
                for p in cell.paragraphs:
                    for run in p.runs:
                        run.bold = True
                        run.font.color.rgb = INK


def add_table(doc, headers, rows, widths):
    table = doc.add_table(rows=1, cols=len(headers))
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            cells[i].text = str(value)
    style_table(table, widths)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return table


def add_kv_table(doc, rows, widths):
    table = doc.add_table(rows=0, cols=2)
    for label, value in rows:
        cells = table.add_row().cells
        cells[0].text = label
        cells[1].text = value
    style_table(table, widths, header=False)
    for row in table.rows:
        set_cell_shading(row.cells[0], LIGHT_GRAY)
        for run in row.cells[0].paragraphs[0].runs:
            run.bold = True
            run.font.color.rgb = INK
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return table


def add_link(paragraph, text, url):
    part = paragraph.part
    rel_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), rel_id)
    run = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "0563C1")
    r_pr.append(color)
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    r_pr.append(underline)
    run.append(r_pr)
    text_node = OxmlElement("w:t")
    text_node.text = text
    run.append(text_node)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def section_competition_bar(doc):
    add_heading(doc, "1. Competition Bar and Business Target")
    add_para(
        doc,
        "VIDSLOOM must be treated as a business launch, not a prototype submission. The XPRIZE page requires a real business with real customers and revenue; the XPRIZE announcement describes the judging criteria as business viability, AI-native operations, and category impact.",
    )
    add_table(
        doc,
        ["Signal", "VIDSLOOM interpretation", "Deadline evidence"],
        [
            (
                "Business viability",
                "Paid customer workflow with repeatable pricing, conversion motion, retention loop, and expense disclosure.",
                "Stripe/invoice exports, customer list, P&L, CAC/spend log, case studies.",
            ),
            (
                "AI-native operations",
                "Gemini agents perform material operating work: Search-grounded trend intelligence, scripts, creative QA, publishing queue preparation, campaign recommendations, customer reporting.",
                "Agent logs, model/API usage records, prompt/output archive, Cloud Logging screenshots.",
            ),
            (
                "Category impact",
                "Small businesses get a practical marketing capability they could not staff themselves.",
                "Before/after customer metrics, testimonials, customer interviews, published campaign examples.",
            ),
        ],
        [1.35, 2.65, 2.5],
    )
    add_heading(doc, "Recommended Deadline Targets", 2)
    add_table(
        doc,
        ["Target level", "By August 17", "Use in submission"],
        [
            ("Minimum credible", "5 paying customers; at least $1,000 collected; 3 case studies.", "Proves real revenue and working service delivery."),
            ("Strong", "15 paying customers; $5,000+ collected or contracted; $2,000+ MRR.", "Shows repeatable demand and early product-market pull."),
            ("Stretch", "30 paying customers or 5 agency clients; $10,000+ collected; 10 public testimonials.", "Creates a standout business-viability story."),
        ],
        [1.35, 2.8, 2.35],
    )
    add_para(
        doc,
        "Use collected revenue rather than forecasted revenue wherever possible. If revenue is still developing, separate paid invoices, signed commitments, and active trials so judges can see what is real.",
    )


def section_scope(doc):
    add_heading(doc, "2. Product Scope: What Must Exist by Deadline")
    add_para(
        doc,
        "The original business plan describes a broad autonomous platform across TikTok, Instagram, YouTube, X, Facebook, and LinkedIn. That is the long-term product. The deadline product should be a narrower AI content operations system that can be sold and proven quickly.",
    )
    add_table(
        doc,
        ["Capability", "Hackathon MVP", "Post-deadline expansion"],
        [
            (
                "Customer onboarding",
                "Business profile, offer, audience, brand voice, asset notes, approval policy, notification channels, social account readiness, budget preference.",
                "CRM integrations, team roles, multi-brand agency accounts.",
            ),
            (
                "Trend intelligence",
                "ZeitgeistScout uses Gemini with Google Search grounding to capture current short-form trend patterns, then maps them to customer-specific remix formulas.",
                "Continuous cross-platform crawling, velocity prediction, competitor benchmarking.",
            ),
            (
                "Content generation",
                "Gemini scripts, hooks, captions, shot lists, image prompts, basic generated assets, editable review queue.",
                "Veo-driven full video generation, automatic remixing, voice cloning where consented.",
            ),
            (
                "Publishing",
                "OAuth-aware publishing queue with captions, hashtags, schedule windows, asset checks, approval checks, notification rules, and manual-upload fallback.",
                "Direct multi-platform posting, comment workflows, and analytics once each platform grants customer-authorized posting permissions.",
            ),
            (
                "Analytics",
                "Manual import or connected metrics, per-campaign summary, Gemini recommendations, revenue notes.",
                "Automated attribution, conversion tracking, cohort retention analytics.",
            ),
            (
                "Evidence logging",
                "Every AI run logs model, prompt version, inputs, outputs, approval, and customer result.",
                "Audit dashboard, agent replay, customer-facing proof reports.",
            ),
        ],
        [1.35, 2.65, 2.5],
    )
    add_heading(doc, "Non-Negotiable MVP Release Criteria", 2)
    add_bullets(
        doc,
        [
            "A customer can pay through Stripe, invoice, or payment link before receiving service.",
            "A customer can complete onboarding without a founder interview, even if the founder still reviews quality.",
            "VIDSLOOM can produce at least one weekly campaign pack per customer with Gemini-generated strategy, copy, and creative direction.",
            "Each campaign has a logged AI decision trail that can be shown in the demo video.",
            "Every paid customer has a permissioned evidence record: contact, problem, deliverable, result, and testimonial request status.",
        ],
    )


def section_technical(doc):
    add_heading(doc, "3. Technical Implementation")
    add_para(
        doc,
        "Build the system as a managed Google Cloud product from day one, but keep the first version simple enough to ship in weeks. The architecture should optimize for evidence, traceability, and fast customer delivery rather than premature scale.",
    )
    add_table(
        doc,
        ["Layer", "Recommended implementation", "Reason"],
        [
            ("Frontend", "Next.js public site at / and campaign workbench at /app, deployed to Cloud Run.", "Clear customer-facing surface plus a usable demo/operator workflow."),
            ("API", "TypeScript or Python service on Cloud Run with OpenAPI documentation.", "Managed deployment and clear contract for agents."),
            ("Database", "Firestore for users, campaigns, agent runs, customer accounts, and evidence index.", "Low-ops, real-time-friendly, flexible schema."),
            ("Files", "Cloud Storage buckets for uploaded assets, generated videos, screenshots, invoices, and evidence files.", "Durable object storage with signed URLs."),
            ("Events", "Pub/Sub topics for campaign generation, reporting, and evidence capture jobs.", "Async workflow without brittle long-running requests."),
            ("Scheduling", "Cloud Scheduler triggers for trend scans, weekly campaign generation, and report emails.", "Visible AI-native operations cadence."),
            ("AI", "Gemini through Vertex AI; Search-grounded ZeitgeistScout separated from controlled JSON campaign generation; Imagen/Veo where available and practical.", "Core XPRIZE alignment with Google AI stack while preserving reliable structured outputs."),
            ("Observability", "Cloud Logging, Error Reporting, and simple admin dashboard for agent status.", "Submission-ready proof of production operation."),
            ("Payments", "Stripe payment links/subscriptions plus invoice backup.", "Fastest route to revenue evidence."),
        ],
        [1.1, 3.0, 2.4],
    )
    add_heading(doc, "Core Data Objects", 2)
    add_table(
        doc,
        ["Object", "Fields to capture first", "Evidence value"],
        [
            ("Customer", "name, business type, contact, consent status, plan, payment status, acquisition source.", "Supports customer proof and CAC analysis."),
            ("Brand profile", "offer, audience, tone, prohibited claims, approved examples, assets.", "Shows personalization and compliance controls."),
            ("Campaign", "goal, target platform, trend reference, hook, script, captions, publishing status, schedule window, approval state.", "Demo-ready product workflow."),
            ("Agent run", "agent name, model, prompt version, input hash, output, tool calls, reviewer, timestamp.", "AI-native operations proof."),
            ("Metric snapshot", "views, clicks, leads, sales, customer-reported result, source, capture date.", "Business impact evidence."),
            ("Expense record", "category, amount, vendor, date, receipt link, notes.", "Submission expense disclosure."),
        ],
        [1.3, 3.0, 2.2],
    )
    add_heading(doc, "Agent Architecture", 2)
    add_table(
        doc,
        ["Agent", "First job", "Human control point"],
        [
            ("OfferProfiler", "Turn onboarding answers into a structured brand and offer brief.", "Customer confirms or edits claims before use."),
            ("ZeitgeistScout", "Use Gemini with Search grounding to identify current short-form trend formulas and freshness caveats.", "Trend outputs are time-bounded and never claim guaranteed virality."),
            ("TrendScout", "Score trend ideas against the customer's audience, offer, cost sensitivity, and available assets.", "Founder/customer approves formulas before high-spend use."),
            ("ScriptForge", "Generate hooks, scripts, captions, hashtags, and CTAs.", "Approval queue before customer delivery."),
            ("CreativeDirector", "Produce shot lists, visual prompts, editing notes, thumbnails, and video assembly instructions.", "Review for brand safety and platform fit."),
            ("PublisherAssist", "Prepare OAuth-aware publishing queue, captions, hashtags, schedules, asset checks, and notification rules.", "Manual publish or direct API only after customer account permissions are connected."),
            ("RevenueAnalyst", "Summarize performance and recommend next campaign changes.", "Customer-facing weekly report reviewed before send."),
        ],
        [1.35, 3.0, 2.15],
    )
    add_heading(doc, "Build Order", 2)
    add_numbered(
        doc,
        [
            "Create the public site, /app workbench, authentication, customer model, and billing links.",
            "Build onboarding, automation settings, brand profile capture, and notification preferences before deeper creative automation.",
            "Implement Gemini campaign generation with Search-grounded ZeitgeistScout, prompt versioning, timeout controls, and output logging.",
            "Add OAuth-aware approval queue, campaign export, and evidence capture.",
            "Add customer asset uploads, metrics import, and weekly AI-generated reporting.",
            "Only then add direct platform posting and fully automated scheduling for connected/approved customer accounts.",
        ],
    )


def section_revenue(doc):
    add_heading(doc, "4. Revenue and Customer Acquisition")
    add_para(
        doc,
        "The fastest path to revenue is a service-assisted SaaS wedge: customers pay for a weekly AI-generated campaign engine, while the team manually closes any automation gaps behind the scenes. The workflow still operates with AI; humans are supervising, selling, and ensuring quality.",
    )
    add_table(
        doc,
        ["Offer", "Price", "What customer gets", "Why it works by deadline"],
        [
            (
                "Founding Launch Sprint",
                "$299 one-time or $149/mo beta",
                "14 days of AI campaign packs, 6-10 short-form concepts, captions, publishing checklist, weekly report.",
                "Low-friction first payment and fast proof.",
            ),
            (
                "VIDSLOOM Ignite Beta",
                "$149/mo",
                "Weekly content engine, 20 video scripts/month, trend briefs, approved posting kits, analytics review.",
                "Matches current business plan while staying deliverable.",
            ),
            (
                "Agency Pilot",
                "$399/mo",
                "White-label workflow for 3 client brands, shared campaign dashboard, priority generation.",
                "Higher revenue with fewer sales if agency relationships exist.",
            ),
            (
                "Revenue share add-on",
                "3-5% attributable uplift",
                "Used only where conversion tracking or customer reporting is credible.",
                "Aligns upside but should not block initial paid subscriptions.",
            ),
        ],
        [1.2, 1.0, 2.55, 1.75],
    )
    add_heading(doc, "Sales Motion", 2)
    add_table(
        doc,
        ["Stage", "Target volume per week", "Definition of done"],
        [
            ("Prospect list", "100 named SMBs/agencies", "Business type, owner/contact, social links, likely pain, outreach channel."),
            ("Outreach", "60 direct messages/emails", "Personalized message with one generated content idea or trend angle."),
            ("Discovery", "10 calls or async audits", "Pain, budget, current content cadence, decision maker, permission to create sample."),
            ("Paid conversion", "3-5 payments", "Stripe/invoice paid or signed commitment with first delivery date."),
            ("Delivery", "All paid customers served weekly", "Campaign pack delivered, feedback requested, metrics capture scheduled."),
        ],
        [1.35, 1.65, 3.5],
    )
    add_heading(doc, "Customer Wedge", 2)
    add_para(
        doc,
        "Prioritize customers with an existing offer and clear conversion path. The best first customers are local service businesses, Shopify/D2C brands, coaches, course creators, and small agencies that already believe social video matters but cannot produce consistently.",
    )
    add_bullets(
        doc,
        [
            "Avoid customers who need a full marketing strategy from zero; choose businesses with active offers.",
            "Avoid regulated claims-heavy niches at first unless the team can review compliance.",
            "Prefer customers who can provide measurable results within two weeks: bookings, leads, sales, traffic, or follower growth.",
            "Ask every paying customer for permission to include name, contact, result summary, and testimonial in XPRIZE evidence.",
        ],
    )


def section_evidence(doc):
    add_heading(doc, "5. Evidence and Operating Cadence")
    add_para(
        doc,
        "Evidence capture is a product requirement. It should happen during normal operations so the final submission is assembled from live records rather than manually reconstructed proof.",
    )
    add_table(
        doc,
        ["Evidence type", "Capture method", "Owner/cadence"],
        [
            ("Revenue", "Stripe export, paid invoice PDFs, bank/payment screenshots, monthly P&L.", "Finance owner; update every Friday."),
            ("Expenses", "Cloud billing, AI/API usage, marketing spend, contractors, zero-spend rows where applicable.", "Ops owner; update every Friday."),
            ("Customers", "Customer roster, consent status, testimonial request, before/after notes, contact details.", "Sales owner; update after each onboarding."),
            ("Product", "Agent logs, campaign screenshots, Cloud Logging screenshots, API usage, dashboard captures.", "Engineering owner; capture continuously."),
            ("Impact", "Leads, bookings, sales, traffic, views, engagement, quotes from customers.", "Customer success owner; update weekly."),
        ],
        [1.25, 3.1, 2.15],
    )
    add_heading(doc, "Weekly Operating Review", 2)
    add_table(
        doc,
        ["Metric", "Target by week 4", "Target by deadline"],
        [
            ("Paying customers", "5", "15 strong / 30 stretch"),
            ("Collected revenue", "$1,000+", "$5,000 strong / $10,000 stretch"),
            ("Customer campaign packs delivered", "15 cumulative", "100+ cumulative"),
            ("Agent runs logged", "100+", "1,000+"),
            ("Case studies", "1 complete", "5+ complete"),
            ("Testimonials", "3 draft quotes", "10 permissioned quotes"),
            ("Marketing/CAC spend log", "Complete", "Complete and reconciled"),
        ],
        [2.35, 1.85, 2.3],
    )
    add_heading(doc, "Daily Founder Cadence", 2)
    add_numbered(
        doc,
        [
            "Morning: review agent failures, customer blockers, and active deliverables.",
            "Midday: do outreach and sales follow-up before product tinkering.",
            "Afternoon: ship the smallest product improvement that removes one delivery bottleneck.",
            "Evening: capture evidence, update metrics, and generate next-day task list with Gemini.",
        ],
    )


def section_timeline(doc):
    add_heading(doc, "6. Start-to-Finish Timeline")
    add_para(
        doc,
        "This timeline assumes work starts immediately on June 13, 2026 and treats August 15 as the practical submission freeze, leaving August 16-17 for corrections only.",
    )
    rows = [
        ("Sprint 0", "Jun 13-15", "Workspace, offer, evidence system", "Repo organized, public site, /app workbench, Cloud Run staging/production, evidence folders, prospect list v1.", "Draft launch offer and start outreach."),
        ("Sprint 1", "Jun 16-22", "Sellable automation MVP", "Automation setup, Search-grounded trend agent, Gemini campaign generator, publishing queue, agent-run log, first posting-kit template.", "Collect first payments from warm contacts."),
        ("Sprint 2", "Jun 23-29", "First customer delivery", "Deliver campaign packs to first customers; capture testimonials and iteration notes.", "5 paying customers or signed commitments."),
        ("Sprint 3", "Jun 30-Jul 6", "Dashboard alpha", "Customer dashboard, approval queue, notification delivery, Firestore data model, storage buckets.", "Publish first case study draft."),
        ("Sprint 4", "Jul 7-13", "Paid beta", "Stripe subscriptions, weekly reports, metric snapshots, product evidence dashboard.", "$1,000+ collected; 100+ agent runs."),
        ("Sprint 5", "Jul 14-20", "Automation depth", "OAuth account connections where approved, Pub/Sub jobs, Scheduler-triggered campaign generation, RevenueAnalyst reports, error logging.", "10 paying customers; repeatable weekly operations."),
        ("Sprint 6", "Jul 21-27", "Scale outreach", "Agency pilot packaging, referral offer, template library, customer onboarding refinements.", "$2,500+ collected; 3 case studies."),
        ("Sprint 7", "Jul 28-Aug 3", "Submission story proof", "Demo script, video storyboard, evidence index, public/private repo readiness.", "15 paying customers or equivalent contracted revenue."),
        ("Sprint 8", "Aug 4-10", "Product freeze", "Bug fixes, reliability pass, security review, final screenshots, narrative draft.", "All required evidence types present."),
        ("Sprint 9", "Aug 11-17", "Submission", "3-minute video, final narrative, revenue/expense docs, GitHub sharing, Devpost final check.", "Submit by Aug 15 if possible; keep Aug 16-17 for edits."),
    ]
    add_table(
        doc,
        ["Sprint", "Dates", "Theme", "Build deliverables", "Business target"],
        rows,
        [0.75, 0.85, 1.2, 2.45, 1.25],
    )
    add_heading(doc, "Final 10-Day Submission Runbook", 2)
    add_table(
        doc,
        ["Date window", "Action"],
        [
            ("Aug 7-9", "Record raw product demo clips: onboarding, agent logs, campaign generation, dashboard, customer proof."),
            ("Aug 10-11", "Lock written narrative at 500-1000 words and reconcile every claim with evidence."),
            ("Aug 12", "Export revenue and expense files; redact private data where needed."),
            ("Aug 13", "Share GitHub repo with required judging/testing accounts and verify access."),
            ("Aug 14", "Render final three-minute video and upload to a stable URL."),
            ("Aug 15", "Complete Devpost submission and save/submit early."),
            ("Aug 16-17", "Only fix broken links, permissions, typos, or evidence gaps. Do not add major new scope."),
        ],
        [1.25, 5.25],
    )


def section_risk(doc):
    add_heading(doc, "7. Risk Register")
    add_table(
        doc,
        ["Risk", "Impact", "Mitigation"],
        [
            ("Social platform API approvals take too long.", "Direct auto-posting blocked.", "Use OAuth-aware posting queues, manual upload fallback, platform-native scheduling where available, and record AI workflow proof."),
            ("Search-grounded trend refresh is slow or unavailable.", "Campaign generation feels slow or trend layer falls back.", "Use bounded Gemini timeouts, freshness caveats, deterministic fallback formulas, and scheduled/background trend refreshes."),
            ("Customers will not pay for an unfinished dashboard.", "Revenue target missed.", "Sell the outcome as a founding launch sprint with clear deliverables, not unfinished software access."),
            ("Generated videos are inconsistent or expensive.", "Quality and gross margin risk.", "Start with scripts, creative briefs, image prompts, templates, and assisted video drafts; reserve full video generation for demos or premium cases."),
            ("Revenue evidence is weak.", "Business viability score suffers.", "Collect payment upfront, record invoices, and separate cash collected from unpaid commitments."),
            ("Claims about autonomy overstate reality.", "Credibility risk.", "Be precise: AI runs core operating tasks; humans supervise quality, sales, customer relationships, and compliance."),
            ("Customer data privacy concerns.", "Trust and submission risk.", "Get written permission for testimonials and redact contact/financial data in public materials."),
            ("Too much time spent polishing docs or pitch.", "Product/revenue slips.", "Freeze docs weekly; product and sales are the source of evidence."),
        ],
        [2.0, 1.45, 3.05],
    )


def section_submission(doc):
    add_heading(doc, "8. Devpost Submission Preparation")
    add_para(
        doc,
        "The submission should tell a proof-based story: VIDSLOOM launched a paid AI-operated content engine for small businesses, used Gemini and Google Cloud in production, and created measurable marketing capacity for customers who could not staff it themselves.",
    )
    add_table(
        doc,
        ["Requirement", "Status now", "Deadline-ready artifact"],
        [
            ("GitHub repo", "Local repo contains the Next.js app, Cloud Run deploy script, docs, and evidence structure.", "Public/private GitHub repo shared with testing@devpost.com and judging@hacker.fund."),
            ("3-minute video", "Not started.", "Concise demo showing AI agents, production logs, customer workflow, and business proof."),
            ("Written narrative", "Draft copy created in docs/submission.", "500-1000 word narrative linked tightly to evidence."),
            ("Revenue evidence", "Folder prepared.", "Stripe export, invoices, P&L, and revenue summary."),
            ("Expenses", "Folder prepared.", "Cloud/API/marketing spend disclosure, including zero-spend categories."),
            ("Product evidence", "Production app, API endpoints, Firestore-backed evidence feed, and Gemini smoke campaigns exist.", "Agent logs, API usage records, Cloud screenshots, dashboard screenshots."),
            ("Customer evidence", "Folder prepared.", "Customer list, contact info, testimonials, results, permissions."),
        ],
        [1.55, 2.1, 2.85],
    )
    add_heading(doc, "Narrative Draft Skeleton", 2)
    add_para(
        doc,
        "VIDSLOOM is an AI-native content operations business for small businesses. It uses Gemini-powered agents to transform a customer's offer, audience, brand assets, automation preferences, and social account readiness into weekly short-form video campaigns, then uses performance feedback to improve the next campaign cycle. During the hackathon, the business should focus on paid customers and documented outcomes: customer onboarding, Search-grounded trend intelligence, AI-generated campaign packs, OAuth-aware publishing queues, approval workflows, performance reports, and evidence that the system is operating in production.",
    )
    add_para(
        doc,
        "Humans remain responsible for selling, customer relationships, compliance decisions, social account authorization, and final quality judgment. AI handles the repeatable operating work: trend discovery, content strategy, script drafting, visual direction, caption generation, queue preparation, report drafting, and performance recommendations. This is the right balance for a real small-business service: autonomous enough to create leverage, supervised enough to protect trust.",
    )


def section_sources(doc):
    add_heading(doc, "9. Sources and Baseline Materials")
    add_para(doc, "The implementation plan was prepared from the local VIDSLOOM business plan and current official hackathon materials checked on June 13, 2026.")
    sources = [
        ("Local source", "docs/business/VIDSLOOM_Business_Plan.docx"),
        ("Devpost overview", "https://xprize.devpost.com/"),
        ("Devpost official rules", "https://xprize.devpost.com/rules"),
        ("XPRIZE announcement", "https://www.xprize.org/news/xprize-launches-hackathon-with-2-million-prize-pool-backed-by-google"),
        ("Vertex AI Search grounding", "https://cloud.google.com/vertex-ai/generative-ai/docs/grounding/grounding-with-google-search"),
        ("TikTok Content Posting API", "https://developers.tiktok.com/doc/content-posting-api-get-started/"),
        ("YouTube video insert API", "https://developers.google.com/youtube/v3/docs/videos/insert"),
        ("Meta Instagram content publishing", "https://developers.facebook.com/docs/instagram-platform/content-publishing/"),
    ]
    table = doc.add_table(rows=1, cols=2)
    table.rows[0].cells[0].text = "Source"
    table.rows[0].cells[1].text = "Location"
    for label, url in sources:
        cells = table.add_row().cells
        cells[0].text = label
        p = cells[1].paragraphs[0]
        if url.startswith("http"):
            add_link(p, url, url)
        else:
            p.text = url
    style_table(table, [1.55, 4.95])


def build():
    doc = Document()
    configure_section(doc.sections[0])
    configure_styles(doc)
    add_title_page(doc)
    section_competition_bar(doc)
    section_scope(doc)
    section_technical(doc)
    section_revenue(doc)
    section_evidence(doc)
    section_timeline(doc)
    section_risk(doc)
    section_submission(doc)
    section_sources(doc)
    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build()
